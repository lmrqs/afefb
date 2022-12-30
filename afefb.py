import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QLineEdit, QRadioButton, QPushButton, QProgressBar, QFileDialog, QMessageBox, QCheckBox, QListWidget
import facebook
import requests
import re
import docx

class CommentExtractor(QMainWindow):
    def __init__(self):
        super().__init__()

        # Set window title and size
        self.setWindowTitle("Facebook Comment Extractor")
        self.resize(500, 400)

        # Create input field for post ID
        self.post_id_label = QLabel(self)
        self.post_id_label.setText("Post ID(s):")
        self.post_id_label.move(20, 20)
        self.post_id_entry = QLineEdit(self)
        self.post_id_entry.move(20, 50)

        # Create checkbox for single/multiple post selection
        self.multiple_posts_checkbox = QCheckBox(self)
        self.multiple_posts_checkbox.setText("Extract comments from multiple posts")
        self.multiple_posts_checkbox.move(20, 90)

        # Create radio buttons for output format
        self.format_label = QLabel(self)
        self.format_label.setText("Output format:")
        self.format_label.move(20, 130)
        self.word_button = QRadioButton(self)
        self.word_button.setText("Word")
        self.word_button.move(20, 160)
        self.pdf_button = QRadioButton(self)
        self.pdf_button.setText("PDF")
        self.pdf_button.move(20, 190)

        # Create checkbox for single/multiple file selection
        self.multiple_files_checkbox = QCheckBox(self)
        self.multiple_files_checkbox.setText("Create a separate file for each post")
        self.multiple_files_checkbox.move(20, 230)

        # Create output file selection button
        self.output_button = QPushButton(self)
        self.output_button.setText("Select output file")
        self.output_button.move(20, 270)
        self.output_button.clicked.connect(self.select_output)

        # Create list widget to display comments
        self.comments_list = QListWidget(self)
        self.comments_list.move(300, 20)

        # Create extract button
        self.extract_button = QPushButton(self)
        self.extract_button.setText("Extract comments")
        self.extract_button.move(20, 300)
        self.extract_button.clicked.connect(self.extract)
        self.progress_bar = QProgressBar(self)
        self.progress_bar.move(20, 340)
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)

    def select_output(self):
        """
        Opens a file selection dialog for the user to choose the output file.
        """
        self.output_file, _ = QFileDialog.getSaveFileName(self, "Select output file", "", "Word Document (*.docx);;PDF (*.pdf)")

    def extract(self):
        """
        Extracts comments from the specified Facebook post(s) and creates a Word or PDF document with the extracted comments.
        """
        # Get post ID(s) and output format
        post_ids = self.post_id_entry.text().split(",")
        output_format = self.word_button.isChecked()

        # Validate input
        if not post_ids:
            # Display error message if no post ID(s) are entered
            QMessageBox.warning(self, "Error", "Please enter a post ID.")
            return
        if not self.output_file:
            # Display error message if no output file is selected
            QMessageBox.warning(self, "Error", "Please select an output file.")
            return

        # Extract comments from Facebook posts
        comments = []
        for post_id in post_ids:
            try:
                # Extract comments and reactions
                post_comments = self.extract_comments(post_id, like_icon=self.get_like_icon())
                post_reactions = self.extract_reactions(post_id)
            except Exception as e:
                # Display error message if comments or reactions could not be extracted
                QMessageBox.warning(self, "Error", f"Could not extract comments or reactions for post {post_id}: {e}")
                continue

            # Add comments and reactions to list
            comments += post_comments
            self.reactions.update(post_reactions)

        # Create output file
        if output_format:
            # Create Word document
            document = docx.Document()
            for comment in comments:
                document.add_paragraph(comment)
            document.save(self.output_file)
        else:
            # Create PDF document
            pass

    def extract_comments(self, post_id, like_icon):
        """
        Extracts comments from the specified Facebook post and returns them as a list of strings.
        """
        # Get access token
        access_token = self.get_access_token()
        if not access_token:
            return []

        # Extract comments
        url = f"https://graph.facebook.com/v7.0/{post_id}/comments?limit=100&access_token={access_token}"
        comments = []
        while True:
            # Send request to Facebook Graph API
            response = requests.get(url)
            data = response.json()
            for comment in data["data"]:
                # Extract comment text and author
                text = comment["message"].replace("\n", " ")
                author = comment["from"]["name"]

                # Extract reactions for comment
                reactions = self.extract_reactions(comment["id"])

                # Format comment and reactions
                formatted_text = f"{author}: {text}\n"
                for reaction_type, users in reactions.items():
                    if reaction_type == "LIKE":
                        formatted_text += f"{like_icon} {len(users)} "
                    else:
                        formatted_text += f"{reaction_type} {len(users)} "
                    formatted_text += ",".join([user["name"] for user in users[:3]])
                    if len(users) > 3:
                        formatted_text += " +" + str(len(users) - 3)
                    formatted_text += "\n"

                comments.append(formatted_text)

            # Check if there are more comments to retrieve
            if "paging" in data and "next" in data["paging"]:
                url = data["paging"]["next"]
            else:
                break

        return comments

    def extract_reactions(self, post_id):
        """
        Extracts reactions from the specified Facebook post or comment and returns them as a dictionary mapping reaction types to lists of users.
        """
        # Get access token
        access_token = self.get_access_token()
        if not access_token:
            return {}

        # Extract reactions
        reactions = {}
        url = f"https://graph.facebook.com/v7.0/{post_id}/reactions?limit=100&access_token={access_token}"
        while True:
            # Send request to Facebook Graph API
            response = requests.get(url)
            data = response.json()
            for reaction in data["data"]:
                # Add reaction to dictionary
                reaction_type = reaction["type"]
                if reaction_type not in reactions:
                    reactions[reaction_type] = []
                reactions[reaction_type].append(reaction["profile"])

            # Check if there are more reactions to retrieve
            if "paging" in data and "next" in data["paging"]:
                url = data["paging"]["next"]
            else:
                break

        return reactions

     def get_access_token(self):
        """
        Returns a valid Facebook Graph API access token.
        """
        # Check if access token has expired
        if self.access_token and self.token_expires_at > time.time():
            return self.access_token

        # Get access token from Facebook Graph API
        app_id = "your_app_id"
        app_secret = "your_app_secret"
        access_token_url = f"https://graph.facebook.com/v7.0/oauth/access_token?client_id={app_id}&client_secret={app_secret}&grant_type=client_credentials"
        response = requests.get(access_token_url)
        data = response.json()
        self.access_token = data["access_token"]
        self.token_expires_at = time.time() + data["expires_in"]

        return self.access_token

    def create_document(self, comments, format):
        """
        Creates a Word or PDF document with the specified comments.
        """
        # Check if multiple files are requested
        if self.multiple_files_checkbox.isChecked():
            # Create a separate file for each set of comments
            for i, comment_set in enumerate(comments):
                # Create output file name
                file_name, file_ext = os.path.splitext(self.output_file)
                output_file = f"{file_name} ({i+1}){file_ext}"

                # Create document
                self.create_single_document(comment_set, output_file, format)
        else:
            # Create a single file with all comments
            self.create_single_document(comments, self.output_file, format)

    def create_single_document(self, comments, output_file, format):
        """
        Creates a single Word or PDF document with the specified comments.
        """
        if format == "Word":
            # Create Word document
            document = docx.Document()
            for comment in comments:
                # Add comment to document
                document.add_paragraph(comment)
            # Add reactions table to document
            document.add_table(self.create_reactions_table())
            # Save document
            document.save(output_file)
        else:
            # Create PDF document
            pass

    def get_like_icon(self):
        """
        Returns the Unicode character for the like icon.
        """
        return "\u2764"

    def create_reactions_table(self):
        """
        Creates a table with the reactions on top and the users who have that reaction (photo and username) in the respective column.
        """
        # Create table
        table = docx.Table(rows=len(self.reactions) + 1, cols=2)

        # Add reactions to top row
        table.cell(0, 0).text = "Reaction"
        table.cell(0, 1).text = "Users"
        row = 1
        for reaction, users in self.reactions.items():
            # Add reaction to table
            table.cell(row, 0).text = reaction
            # Add user photos and names to table
            for user in users:
                photo_data = requests.get(user["picture"]["data"]["url"]).content
                photo = BytesIO(photo_data)
                table.cell(row, 1).add_paragraph().add_run().add_picture(photo)
                table.cell(row, 1).add_paragraph(user["name"])
                row += 1

        return table

    def create_single_document(self, comments, output_file, format):
        """
        Creates a single Word or PDF document with the specified comments.
        """
        if format == "Word":
            # Create Word document
            document = docx.Document()
            for comment in comments:
                # Add comment to document
                document.add_paragraph(comment)
            # Add reactions table to document
            document.add_table(self.create_reactions_table())
            # Save document
            document.save(output_file)
        else:
            # Create PDF document
            pdf = FPDF()
            # Add comments to PDF
            for comment in comments:
                pdf.add_page()
                pdf.set_xy(10, 10)
                pdf.set_font("Arial", size=12)
                pdf.write(5, comment)
            # Add reactions table to PDF
            self.add_reactions_table_to_pdf(pdf)
            # Save PDF
            pdf.output(output_file)

    def add_reactions_table_to_pdf(self, pdf):
        """
        Adds a table with the reactions on top and the users who have that reaction (photo and username) in the respective column to the specified PDF document.
        """
        # Calculate table dimensions
        num_rows = len(self.reactions) + 1
        num_cols = 2
        cell_width = 60
        cell_height = 20
        table_width = cell_width * num_cols
        table_height = cell_height * num_rows

        # Add table to PDF
        pdf.set_xy(10, 10)
        pdf.cell(table_width, table_height, "", 1)
        pdf.set_xy(10, 10)
        pdf.cell(cell_width, cell_height, "Reaction", 1)
        pdf.cell(cell_width, cell_height, "Users", 1)
        row = 1
        for reaction, users in self.reactions.items():
            # Add reaction to table
            pdf.set_xy(10, 10 + cell_height * (row + 1))
            pdf.cell(cell_width, cell_height, reaction, 1)
            # Add user photos and names to table
            for user in users:
                # Download user photo
                photo_data = requests.get(user["picture"]["data"]["url"]).content
                photo = Image.open(BytesIO(photo_data))
                # Resize photo to fit in cell
                photo.thumbnail((cell_width, cell_height))
                # Resize photo to fit in cell
                photo.thumbnail((cell_width, cell_height))
                # Add photo to PDF
                pdf.image(image=BytesIO(photo_data), x=10 + cell_width * (row + 1), y=10 + cell_height * (col + 1), w=cell_width, h=cell_height)
                # Add user name to table
                pdf.set_xy(10 + cell_width * (row + 1), 10 + cell_height * (col + 1))
                pdf.cell(cell_width, cell_height, user["name"], 0, 0, "C")
                col += 1
            row += 1
    def create_multiple_documents(self, comments_by_post, output_files, format):
        """
        Creates a separate Word or PDF document for each post with the respective comments.
        """
        for post_id, comments in comments_by_post.items():
            # Create document for post
            self.create_single_document(comments, output_files[post_id], format)

    def create_reactions_table(self):
        """
        Creates a table with the reactions on top and the users who have that reaction (photo and username) in the respective column.
        """
        # Create table
        table = docx.Table(rows=len(self.reactions) + 1, cols=2)
        # Add reactions to table
        table.cell(0, 0).text = "Reaction"
        table.cell(0, 1).text = "Users"
        row = 1
        for reaction, users in self.reactions.items():
            # Add reaction to table
            table.cell(row, 0).text = reaction
            # Add user photos and names to table
            for user in users:
                # Download user photo
                photo_data = requests.get(user["picture"]["data"]["url"]).content
                photo = Image.open(BytesIO(photo_data))
                # Add photo to table
                table.cell(row, 1).add_paragraph().add_run().add_picture(BytesIO(photo_data), width=Inches(0.5))
                # Add user name to table
                table.cell(row, 1).add_paragraph(user["name"])
                row += 1
        return table

    def create_single_document(self, comments, output_file, format):
        """
        Creates a Word or PDF document with the specified comments.
        """
        if format == "Word":
            # Create Word document
            document = docx.Document()
            # Add comments to document
            for comment in comments:
                # Add comment text to document
                document.add_paragraph(comment["message"])
                # Add comment author and date to document
                document.add_paragraph(f"{comment['from']['name']} - {comment['created_time']}")
            # Add reactions table to document
            if self.reactions:
                document.add_page_break()
                document.add_heading("Reactions", 2)
                document.add_table(self.create_reactions_table())
            # Save document
            document.save(output_file)
        elif format == "PDF":
            # Create PDF document
            pdf = FPDF()
            # Add comments to PDF
            for comment in comments:
                # Add comment text to PDF
                pdf.add_page()
                pdf.set_xy(10, 10)
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, comment["message"])
                # Add comment author and date to PDF
                pdf.set_xy(10, pdf.get_y() + 10)
                pdf.set_font("Arial", size=10)
                pdf.cell(0, 10, f"{comment['from']['name']} - {comment['created_time']}")
            # Add reactions table to PDF
            if self.reactions:
                # Calculate cell size
                num_users = sum(len(users) for users in self.reactions.values())
                cell_width = round(210 / (num_users + 1))
                cell_height = round(297 / (len(self.reactions) + 1))
                # Add table header
                pdf.add_page()
                pdf.set_xy(10, 10)
                pdf.set_font("Arial", "B", size=12)
                pdf.cell(cell_width, cell_height, "Reaction", 1, 0, "C")
                pdf.cell(cell_width * num_users, cell_height, "Users", 1, 0, "C")
                # Add reactions and users to table
                row = 0
                for reaction, users in self.reactions.items():
                    col = 0
                    # Add reaction to table
                    pdf.set_xy(10 + cell_width * (col + 1), 10 + cell_height * (row + 1))
                    pdf.set_font("Arial", size=12)
                    pdf.cell(cell_width, cell_height, reaction, 1, 0, "C")
                    col += 1

    def extract(self):
        """
        Extracts comments from the specified Facebook post(s) and creates a Word or PDF document with the extracted comments.
        """
        # Get post ID(s) and output format
        post_ids = self.post_id_entry.text().split(",")
        output_format = self.word_button.isChecked()

        # Validate input
        if not post_ids:
            # Display error message if no post ID(s) are entered
            QMessageBox.warning(self, "Error", "Please enter a post ID.")
            return
        if not self.output_file:
            # Display error message if no output file is selected
            QMessageBox.warning(self, "Error", "Please select an output file.")
            return

        # Extract comments from Facebook posts
        comments = []
        self.reactions = {}
        for post_id in post_ids:
            try:
                comments += self.extract_comments(post_id)
            except Exception as e:
                # Display error message if there was an issue extracting comments from the post
                QMessageBox.warning(self, "Error", f"There was an issue extracting comments from post {post_id}: {e}")
        # Display comments in list widget
        for comment in comments:
            self.comments_list.addItem(f"{comment['from']['name']}: {comment['message']}")

        # Create output file(s)
        if self.multiple_files_checkbox.isChecked():
            # Create separate files for each post
            comments_by_post = self.group_comments_by_post(comments)
            output_files = {post_id: self.output_file.replace(".docx", f"_{post_id}.docx").replace(".pdf", f"_{post_id}.pdf") for post_id in comments_by_post.keys()}
            self.create_multiple_documents(comments_by_post, output_files, output_format)
        else:
            # Create single file for all posts
            self.create_single_document(comments, self.output_file, output_format)

if __name__ == "__main__":
    # Create and run application
    app = QApplication(sys.argv)
    window = CommentExtractor()
    window.show()
    sys.exit(app.exec_())
